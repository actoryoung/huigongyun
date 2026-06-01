"""Word 文档解析占位实现。

当前为占位解析器，声明对 .doc/.docx 的支持，未来将实现段落与表格抽取、
技术约束挖掘及引用/来源信息的保留。
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from .base import ScaffoldFormatParser
from ..models import ProjectDocument


class WordSourceParser(ScaffoldFormatParser):
    """Word 源解析器实现（可选依赖：python-docx）。

    实现说明：若环境中安装了 `python-docx`（导入名 `docx`），此解析器会尝试
    抽取段落文本与表格内容并返回 `parse_status: ok`。如果未安装该依赖，
    则回退到占位实现以保证注册表路由不发生改变。
    """

    input_kind = "word"
    source_format = "word"
    message = "Word 解析实现：尝试使用 python-docx 抽取段落与表格，缺失依赖时回退。"

    def supported_suffixes(self) -> set[str]:
        return {".doc", ".docx"}

    def parse(self, input_path: str) -> ProjectDocument:
        path = Path(input_path)
        try:
            import docx  # type: ignore
        except Exception:
            # python-docx not available; fall back to scaffold behavior
            return super().parse(input_path)

        # Parse .docx with python-docx
        try:
            document = docx.Document(str(path))
        except Exception:
            return ProjectDocument(
                project_name=path.stem or "project",
                files=[str(path)],
                metadata={
                    "input_kind": "word",
                    "parse_status": "error",
                    "source_format": "word",
                    "message": "Failed to open .docx file with python-docx.",
                },
            )

        paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        tables: list[dict[str, Any]] = []
        for table in document.tables:
            rows: list[list[str]] = []
            for r in table.rows:
                cells = [c.text for c in r.cells]
                rows.append(cells)
            tables.append({"row_count": len(rows), "rows": rows})

        return ProjectDocument(
            project_name=path.stem or "project",
            files=[str(path)],
            metadata={
                "input_kind": "word",
                "parse_status": "ok",
                "source_format": "word",
                "paragraph_count": len(paragraphs),
                "paragraphs": paragraphs,
                "table_count": len(tables),
                "tables": tables,
            },
        )